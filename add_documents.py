#!/usr/bin/env python3
"""
DOCUMENT INGESTION SCRIPT
File name: add_documents.py
Add your own PDF, TXT, and DOCX files to the RAG system
"""

import os
import sys
from pathlib import Path
from typing import List, Dict, Any
import chromadb
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

# Import configuration
from config import config

class DocumentIngestion:
    """Handle document ingestion into the vector database"""
    
    def __init__(self):
        self.setup_components()
    
    def setup_components(self):
        """Initialize embedding model and vector store"""
        print("🔧 Setting up document ingestion system...")
        
        # Load embedding model
        self.embedding_model = HuggingFaceEmbeddings(
            model_name=config.model.embedding_model,
            model_kwargs={'device': 'cuda' if config.performance.use_gpu else 'cpu'}
        )
        
        # Setup text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.model.chunk_size,
            chunk_overlap=config.model.chunk_overlap,
            length_function=len,
        )
        
        # Create ChromaDB client
        os.makedirs(config.database.chroma_path, exist_ok=True)
        self.chroma_client = chromadb.PersistentClient(path=config.database.chroma_path)
        
        print("✅ Document ingestion system ready!")
    
    def load_pdf_file(self, file_path: str) -> List[Document]:
        """Load and process a PDF file"""
        try:
            from PyPDF2 import PdfReader
            
            print(f"📄 Processing PDF: {file_path}")
            reader = PdfReader(file_path)
            
            documents = []
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": os.path.basename(file_path),
                            "page": page_num,
                            "file_type": "pdf"
                        }
                    )
                    documents.append(doc)
            
            print(f"✅ Loaded {len(documents)} pages from PDF")
            return documents
            
        except ImportError:
            print("❌ PyPDF2 not installed. Install with: pip install PyPDF2")
            return []
        except Exception as e:
            print(f"❌ Error processing PDF {file_path}: {e}")
            return []
    
    def load_text_file(self, file_path: str) -> List[Document]:
        """Load and process a text file"""
        try:
            print(f"📝 Processing text file: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if content.strip():
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": os.path.basename(file_path),
                        "file_type": "txt"
                    }
                )
                print(f"✅ Loaded text file")
                return [doc]
            else:
                print("⚠️ File is empty")
                return []
                
        except Exception as e:
            print(f"❌ Error processing text file {file_path}: {e}")
            return []
    
    def load_docx_file(self, file_path: str) -> List[Document]:
        """Load and process a DOCX file"""
        try:
            from docx import Document as DocxDocument
            
            print(f"📄 Processing DOCX: {file_path}")
            doc = DocxDocument(file_path)
            
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            full_content = '\n'.join(content)
            
            if full_content.strip():
                document = Document(
                    page_content=full_content,
                    metadata={
                        "source": os.path.basename(file_path),
                        "file_type": "docx"
                    }
                )
                print(f"✅ Loaded DOCX file")
                return [document]
            else:
                print("⚠️ DOCX file is empty")
                return []
                
        except ImportError:
            print("❌ python-docx not installed. Install with: pip install python-docx")
            return []
        except Exception as e:
            print(f"❌ Error processing DOCX {file_path}: {e}")
            return []
    
    def load_single_file(self, file_path: str) -> List[Document]:
        """Load a single file based on its extension"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"❌ File not found: {file_path}")
            return []
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self.load_pdf_file(str(file_path))
        elif extension == '.txt':
            return self.load_text_file(str(file_path))
        elif extension == '.docx':
            return self.load_docx_file(str(file_path))
        elif extension == '.md':
            # Treat markdown as text
            return self.load_text_file(str(file_path))
        else:
            print(f"❌ Unsupported file type: {extension}")
            print("Supported types: .pdf, .txt, .docx, .md")
            return []
    
    def load_directory(self, directory_path: str) -> List[Document]:
        """Load all supported files from a directory"""
        directory = Path(directory_path)
        
        if not directory.exists():
            print(f"❌ Directory not found: {directory_path}")
            return []
        
        supported_extensions = {'.pdf', '.txt', '.docx', '.md'}
        all_documents = []
        
        print(f"📁 Processing directory: {directory_path}")
        
        # Find all supported files
        files_found = []
        for extension in supported_extensions:
            files_found.extend(directory.glob(f"*{extension}"))
            files_found.extend(directory.glob(f"**/*{extension}"))  # Recursive
        
        if not files_found:
            print("❌ No supported files found in directory")
            print("Supported types: .pdf, .txt, .docx, .md")
            return []
        
        print(f"📋 Found {len(files_found)} files to process")
        
        # Process each file
        for file_path in files_found:
            documents = self.load_single_file(str(file_path))
            all_documents.extend(documents)
        
        print(f"✅ Loaded {len(all_documents)} total documents")
        return all_documents
    
    def add_documents_to_vectorstore(self, documents: List[Document], replace_existing: bool = False):
        """Add documents to the vector store"""
        
        if not documents:
            print("❌ No documents to add")
            return
        
        print(f"🔄 Processing {len(documents)} documents...")
        
        # Split documents into chunks
        all_chunks = []
        for doc in documents:
            chunks = self.text_splitter.split_documents([doc])
            all_chunks.extend(chunks)
        
        print(f"📄 Created {len(all_chunks)} text chunks")
        
        try:
            if replace_existing:
                print("🗑️ Clearing existing data...")
                # Delete existing collection
                try:
                    self.chroma_client.delete_collection(config.database.collection_name)
                except:
                    pass  # Collection might not exist
            
            # Create or get collection
            vectorstore = Chroma.from_documents(
                documents=all_chunks,
                embedding=self.embedding_model,
                client=self.chroma_client,
                collection_name=config.database.collection_name,
                persist_directory=config.database.chroma_path
            )
            
            print("✅ Documents successfully added to vector database!")
            print(f"📊 Database location: {config.database.chroma_path}")
            print(f"📋 Collection name: {config.database.collection_name}")
            
        except Exception as e:
            print(f"❌ Error adding documents to vector store: {e}")
    
    def list_existing_documents(self):
        """List documents currently in the vector store"""
        try:
            vectorstore = Chroma(
                client=self.chroma_client,
                collection_name=config.database.collection_name,
                embedding_function=self.embedding_model,
            )
            
            # Get collection info
            collection = self.chroma_client.get_collection(config.database.collection_name)
            count = collection.count()
            
            print(f"📊 Current vector database status:")
            print(f"   Collection: {config.database.collection_name}")
            print(f"   Total chunks: {count}")
            print(f"   Location: {config.database.chroma_path}")
            
            if count > 0:
                # Get a sample of documents to show sources
                sample_results = collection.get(limit=10)
                sources = set()
                
                for metadata in sample_results.get('metadatas', []):
                    if metadata and 'source' in metadata:
                        sources.add(metadata['source'])
                
                if sources:
                    print(f"   Sources found: {', '.join(list(sources)[:5])}")
                    if len(sources) > 5:
                        print(f"   ... and {len(sources) - 5} more")
            
        except Exception as e:
            print(f"📊 No existing data found or error: {e}")

def main():
    """Main function with interactive menu"""
    
    print("\n" + "="*60)
    print("📚 FAST NUCES AI ASSISTANT - DOCUMENT INGESTION")
    print("="*60)
    print("Add your own documents to the RAG system")
    print()
    
    # Check dependencies
    missing_deps = []
    try:
        import PyPDF2
    except ImportError:
        missing_deps.append("PyPDF2 (for PDF files)")
    
    try:
        import docx
    except ImportError:
        missing_deps.append("python-docx (for DOCX files)")
    
    if missing_deps:
        print("⚠️ Optional dependencies missing:")
        for dep in missing_deps:
            print(f"   - {dep}")
        print("Install with: pip install PyPDF2 python-docx")
        print()
    
    # Initialize ingestion system
    ingestion = DocumentIngestion()
    
    # Show current status
    ingestion.list_existing_documents()
    print()
    
    # Interactive menu
    while True:
        print("📋 What would you like to do?")
        print("1. Add a single file")
        print("2. Add all files from a directory")
        print("3. Replace all existing data with new files")
        print("4. View current database status")
        print("5. Exit")
        print()
        
        choice = input("Enter your choice (1-5): ").strip()
        
        if choice == '1':
            file_path = input("📄 Enter file path: ").strip().strip('"\'')
            documents = ingestion.load_single_file(file_path)
            if documents:
                ingestion.add_documents_to_vectorstore(documents)
        
        elif choice == '2':
            dir_path = input("📁 Enter directory path: ").strip().strip('"\'')
            documents = ingestion.load_directory(dir_path)
            if documents:
                replace = input("🔄 Replace existing data? (y/N): ").strip().lower()
                ingestion.add_documents_to_vectorstore(documents, replace_existing=(replace == 'y'))
        
        elif choice == '3':
            dir_path = input("📁 Enter directory path for new data: ").strip().strip('"\'')
            documents = ingestion.load_directory(dir_path)
            if documents:
                confirm = input("⚠️ This will DELETE all existing data. Continue? (y/N): ").strip().lower()
                if confirm == 'y':
                    ingestion.add_documents_to_vectorstore(documents, replace_existing=True)
        
        elif choice == '4':
            ingestion.list_existing_documents()
        
        elif choice == '5':
            break
        
        else:
            print("❌ Invalid choice. Please enter 1-5.")
        
        print("\n" + "-"*40 + "\n")
    
    print("👋 Document ingestion complete!")
    print("🚀 Your documents are now ready for the AI assistant!")

if __name__ == "__main__":
    main()